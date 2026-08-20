from langchain_community.document_loaders import PyPDFLoader, TextLoader,CSVLoader
from docx import Document as DocxDocument
from langchain_core.documents import Document

from langchain_text_splitters import RecursiveCharacterTextSplitter

async def load_and_split(file_path: str) -> list:
    ext = file_path.rsplit('.', 1)[-1].lower()

    if ext == 'pdf':
        loader = PyPDFLoader(file_path)
        documents = await loader.aload()
    elif ext == 'docx':
        # 直接用 python-docx 读取，不依赖 LibreOffice
        doc = DocxDocument(file_path)
        text = '\n'.join(para.text for para in doc.paragraphs)
        documents = [Document(page_content=text, metadata={"source": file_path})]
    elif ext in ('txt', 'md'):
        loader = TextLoader(file_path, encoding="utf-8")
        documents = await loader.aload()
    elif ext == 'csv':
        loader = CSVLoader(file_path)
        documents = await loader.aload()
    elif ext == 'doc':
        raise ValueError("不支持 .doc 格式，请转换为 .docx 后上传")
    else:
        loader = TextLoader(file_path, encoding="utf-8")
        documents = await loader.aload()

    # 3. 分块
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=500, #每个块的最大字符数。常用值 500-1000。块越大包含的信息越多但消耗更多 token，块越小检索更精确但可能丢失上下文
        chunk_overlap=50,# 相邻块之间的重叠字符数，默认 0。设为 chunk_size 的 10%-20% 可以避免在块边界截断关键信息
        separators=["\n\n", "\n", "。", "！", "？", " ", ""], #分隔符列表，按优先级尝试
    )
    return splitter.split_documents(documents)